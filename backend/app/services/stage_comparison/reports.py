"""Экспорт итогового отчёта по сессии в md/html/json + опционально pdf/docx.

Отчёты лежат в comparison/sessions/<sid>/reports/.

Каждый отчёт идентифицируется stable report_id (без расширения), а также
хранится в manifest comparison/sessions/<sid>/reports/reports.json. Имя
файла включает timestamp с миллисекундами и short uuid, чтобы исключить
перезапись при создании двух отчётов в одну секунду.
"""
from __future__ import annotations

import html
import json
import logging
import os
import threading
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from . import alignment as alignment_mod
from . import findings as findings_mod
from . import paths as paths_mod
from . import store as store_mod

logger = logging.getLogger(__name__)

_lock = threading.RLock()


SUPPORTED_FORMATS = ("md", "html", "json", "pdf", "docx")
MAX_IMAGE_WIDTH_PT = 240  # PDF
MAX_IMAGE_WIDTH_CM = 8.0  # DOCX


def _utc_now() -> str:
    return datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")


def _new_report_id() -> str:
    """Stable id отчёта: timestamp с ms + short uuid. Используется и как
    suffix файла, и как ключ в manifest."""
    now = datetime.now()
    ts = now.strftime("%Y-%m-%d_%H%M%S")
    ms = f"{now.microsecond // 1000:03d}"
    short = uuid.uuid4().hex[:6]
    return f"{ts}_{ms}_{short}"


def _filename(report_id: str, fmt: str) -> str:
    return f"comparison_report_{report_id}.{fmt}"


def _allocate_filename(session_id: str, fmt: str) -> tuple[str, str, Path]:
    """Вернуть (report_id, filename, full_path) гарантировано-свободный.

    Защита от race: если файл уже существует — генерируем новый suffix."""
    for _ in range(8):
        rid = _new_report_id()
        fname = _filename(rid, fmt)
        out_path = paths_mod.report_path(session_id, fname)
        if not out_path.exists():
            return rid, fname, out_path
    # Маловероятно: коллизия 8 раз подряд → добавим больше энтропии.
    rid = _new_report_id() + "_" + uuid.uuid4().hex[:8]
    fname = _filename(rid, fmt)
    return rid, fname, paths_mod.report_path(session_id, fname)


# ─── Local crop resolution (Задача 2) ────────────────────────────────────

def _resolve_crop_local_path(session_id: str, finding_item: dict, side: str) -> Optional[Path]:
    """Найти локальный crop-файл для side ('left'/'right') в данном finding.

    1. Если в comparison/sessions/<sid>/pairs/<pid>/crops/<side>/ уже есть
       подготовленный crop — отдадим его.
    2. Иначе попробуем render_block_crop(...) — он лениво создаёт PNG.
    3. Если block_id/pair_id отсутствуют — None.

    Только пути внутри comparison/sessions/<sid>/ считаются безопасными
    (защита от path traversal — Задача 2.6).
    """
    if not finding_item:
        return None
    pair_id = finding_item.get("pair_id")
    side_block = (finding_item.get(side) or {}).get("block_id")
    if not pair_id or not side_block:
        return None
    try:
        crops_dir = paths_mod.crops_dir(session_id, pair_id, side).resolve()
    except (KeyError, ValueError, OSError):
        return None
    # 1) Берём первый матчевый PNG (любого target_long_side)
    safe_bid = "".join(c if c.isalnum() else "_" for c in side_block)
    try:
        for f in sorted(crops_dir.glob(f"{safe_bid}_*.png")):
            f_resolved = f.resolve()
            try:
                f_resolved.relative_to(crops_dir)
            except ValueError:
                continue
            if f_resolved.is_file():
                return f_resolved
    except OSError:
        pass
    # 2) Lazy render
    try:
        out = store_mod.render_block_crop(session_id, pair_id, side, side_block)
        if isinstance(out, Path) and out.exists():
            return out.resolve()
    except (KeyError, FileNotFoundError, ValueError, RuntimeError, OSError):
        return None
    return None


def _safe_local_path(session_id: str, p: Optional[Path]) -> Optional[Path]:
    """Path traversal guard: путь должен быть внутри comparison/sessions/<sid>/."""
    if p is None:
        return None
    try:
        session_root = paths_mod.session_dir(session_id).resolve()
        p_resolved = p.resolve()
        p_resolved.relative_to(session_root)
    except (ValueError, OSError):
        return None
    return p_resolved if p_resolved.exists() and p_resolved.is_file() else None


def _filter_items(items: list[dict], filters: dict) -> list[dict]:
    statuses = filters.get("status") or []
    severities = filters.get("severity") or []
    categories = filters.get("category") or []
    types = filters.get("type") or []
    pair_ids = filters.get("pair_id") or []
    include_rejected = bool(filters.get("include_rejected"))
    include_ignored = bool(filters.get("include_ignored"))
    include_resolved = bool(filters.get("include_resolved", True))
    include_child_findings = bool(filters.get("include_child_findings", True))
    out: list[dict] = []
    for it in items:
        if it.get("deleted") and not include_ignored:
            continue
        st = it.get("status")
        if st == "rejected" and not include_rejected:
            continue
        if st == "ignored" and not include_ignored:
            continue
        if st == "resolved" and not include_resolved:
            continue
        if statuses and st not in statuses:
            continue
        if severities and it.get("severity") not in severities:
            continue
        if categories and it.get("category") not in categories:
            continue
        if types and it.get("type") not in types:
            continue
        if pair_ids and it.get("pair_id") not in pair_ids:
            continue
        # Группировка: если есть parent_finding_id и не запрошен включённый
        # вывод детей — пропускаем.
        if not include_child_findings and it.get("parent_finding_id"):
            continue
        out.append(it)
    return out


def _collect_report_data(session_id: str, filters: dict) -> dict:
    session = store_mod.get_session(session_id)
    if session is None:
        raise KeyError("session_not_found")
    findings_data = findings_mod._read_findings(session_id)
    raw_items = findings_data.get("items") or []
    filtered = _filter_items(raw_items, filters)

    pairs = session.get("pairs") or []
    pairs_by_id = {p["id"]: p for p in pairs}

    # Группировка findings по pair_id
    by_pair: dict[str, list[dict]] = {}
    for it in filtered:
        by_pair.setdefault(it.get("pair_id"), []).append(it)

    # Per-pair stats
    pair_blocks: list[dict] = []
    matched_pages_total = 0
    new_right_total = 0
    removed_left_total = 0
    reordered_total = 0
    for p in pairs:
        if p.get("status") == "disabled":
            continue
        pid = p["id"]
        try:
            al = store_mod.get_alignment(session_id, pid)
            stats = alignment_mod.compute_page_stats((al.get("alignment") or {}).get("items") or [])
        except Exception:
            stats = {"matched_pages": 0, "new_right_pages": 0, "removed_left_pages": 0, "reordered_pages": 0}
        matched_pages_total += stats.get("matched_pages", 0)
        new_right_total += stats.get("new_right_pages", 0)
        removed_left_total += stats.get("removed_left_pages", 0)
        reordered_total += stats.get("reordered_pages", 0)
        pair_findings = by_pair.get(pid, [])
        pair_blocks.append({
            "pair": p,
            "stats": stats,
            "findings": pair_findings,
        })

    # Summary
    def _counts(items, key):
        c = {}
        for it in items:
            c[str(it.get(key, "unknown"))] = c.get(str(it.get(key, "unknown")), 0) + 1
        return c

    summary = {
        "total": len(filtered),
        "by_status": _counts(filtered, "status"),
        "by_category": _counts(filtered, "category"),
        "by_severity": _counts(filtered, "severity"),
        "by_type": _counts(filtered, "type"),
        "pages": {
            "matched_pages": matched_pages_total,
            "new_right_pages": new_right_total,
            "removed_left_pages": removed_left_total,
            "reordered_pages": reordered_total,
        },
    }

    # Warnings (Задача 9) — добавляем в отчёт
    warnings_payload = {"items": [], "summary": {"high": 0, "medium": 0, "low": 0}}
    try:
        from . import warnings as warnings_mod
        warnings_payload = warnings_mod.compute_warnings(session_id)
    except Exception:
        logger.exception("compute_warnings failed in report")

    return {
        "session": {
            "id": session.get("id"),
            "created_at": session.get("created_at"),
            "stage_a_path": session.get("stage_a_path"),
            "stage_b_path": session.get("stage_b_path"),
        },
        "generated_at": _utc_now(),
        "filters": filters,
        "summary": summary,
        "pairs": pair_blocks,
        "raw_total": len(raw_items),
        "warnings": warnings_payload.get("items") or [],
        "warnings_summary": warnings_payload.get("summary") or {},
    }


# ─── Renderers ───────────────────────────────────────────────────────────

def _render_markdown(data: dict, *, include_images: bool, include_llm: bool,
                     include_user_notes: bool) -> str:
    s = data["session"]
    lines: list[str] = []
    lines.append(f"# Отчёт по сравнению стадий")
    lines.append("")
    lines.append(f"- Сессия: `{s['id']}`")
    lines.append(f"- Дата формирования: {data['generated_at']}")
    lines.append(f"- Папка первой стадии: `{s.get('stage_a_path')}`")
    lines.append(f"- Папка второй стадии: `{s.get('stage_b_path')}`")
    pp = data["summary"]["pages"]
    pairs_n = len(data["pairs"])
    lines.append(f"- PDF-пар: {pairs_n}")
    lines.append(f"- Сопоставленных листов: {pp['matched_pages']}")
    lines.append(f"- Новых листов справа: {pp['new_right_pages']}")
    lines.append(f"- Удалённых листов слева: {pp['removed_left_pages']}")
    lines.append(f"- С изменённым порядком: {pp['reordered_pages']}")
    lines.append("")

    warnings = data.get("warnings") or []
    if warnings:
        lines.append("## Предупреждения сессии")
        lines.append("")
        for w in warnings:
            lines.append(f"- **[{w.get('severity', 'low')}]** {w.get('title')}"
                         + (f" — {w.get('details')}" if w.get('details') else ""))
        lines.append("")

    su = data["summary"]
    lines.append("## Сводка")
    lines.append("")
    lines.append(f"- Всего расхождений: **{su['total']}**")
    if su["by_status"]:
        lines.append(f"- По статусам: " + ", ".join(f"{k}: **{v}**" for k, v in su["by_status"].items()))
    if su["by_category"]:
        lines.append(f"- По категориям: " + ", ".join(f"{k}: **{v}**" for k, v in su["by_category"].items()))
    if su["by_severity"]:
        lines.append(f"- По важности: " + ", ".join(f"{k}: **{v}**" for k, v in su["by_severity"].items()))
    if su["by_type"]:
        lines.append(f"- По типам: " + ", ".join(f"{k}: **{v}**" for k, v in su["by_type"].items()))
    lines.append("")

    lines.append("## Расхождения по проектам")
    lines.append("")
    finding_global_idx = 0
    for pb in data["pairs"]:
        p = pb["pair"]
        lp = (p.get("left") or {}).get("filename", "—")
        rp = (p.get("right") or {}).get("filename", "—")
        status = p.get("status", "—")
        st = pb["stats"]
        lines.append(f"### {lp} ↔ {rp}")
        lines.append("")
        lines.append(f"- Статус сопоставления: `{status}`")
        lines.append(
            f"- Листы: сопоставлено **{st['matched_pages']}**, "
            f"новых справа **{st['new_right_pages']}**, удалено слева **{st['removed_left_pages']}**, "
            f"переставлено **{st['reordered_pages']}**"
        )
        lines.append(f"- Расхождений в этой паре: **{len(pb['findings'])}**")
        lines.append("")
        for it in pb["findings"]:
            finding_global_idx += 1
            lines.append(f"#### #{finding_global_idx}. {it.get('title') or it.get('type')}")
            lines.append("")
            lines.append(f"- Тип: `{it.get('type')}`")
            lines.append(f"- Категория: `{it.get('category')}`")
            lines.append(f"- Важность: `{it.get('severity')}`")
            lines.append(f"- Статус: `{it.get('status')}`")
            src = it.get("source") or {}
            if src.get("alignment_slot"):
                lines.append(f"- Slot: {src['alignment_slot']}")
            left = it.get("left") or {}
            right = it.get("right") or {}
            if left.get("page") is not None or right.get("page") is not None:
                lines.append(f"- Листы: L={left.get('page', '—')} / R={right.get('page', '—')}")
            if it.get("summary"):
                lines.append(f"- Описание: {it['summary']}")
            if left.get("text"):
                lines.append("")
                lines.append("Старая стадия:")
                lines.append("```")
                lines.append(str(left["text"])[:1000])
                lines.append("```")
            if right.get("text"):
                lines.append("")
                lines.append("Новая стадия:")
                lines.append("```")
                lines.append(str(right["text"])[:1000])
                lines.append("```")
            if include_images and (left.get("crop_url") or right.get("crop_url")):
                lines.append("")
                if left.get("crop_url"):
                    lines.append(f"![left]({left['crop_url']})")
                if right.get("crop_url"):
                    lines.append(f"![right]({right['crop_url']})")
            if include_llm and it.get("llm_summary"):
                lines.append("")
                lines.append("LLM-описание:")
                lines.append("> " + str(it["llm_summary"]).replace("\n", "\n> "))
            if include_user_notes and it.get("user_note"):
                lines.append("")
                lines.append(f"_Комментарий проверяющего:_ {it['user_note']}")
            lines.append("")

    # Приложение
    lines.append("---")
    lines.append("")
    lines.append("## Приложение")
    lines.append("")
    lines.append("### Параметры фильтрации")
    lines.append("```json")
    lines.append(json.dumps(data.get("filters") or {}, ensure_ascii=False, indent=2))
    lines.append("```")
    lines.append("")
    lines.append(f"- Findings всего в сессии (до фильтра): {data['raw_total']}")
    return "\n".join(lines)


def _render_html(data: dict, *, include_images: bool, include_llm: bool,
                 include_user_notes: bool) -> str:
    s = data["session"]
    su = data["summary"]
    pp = su["pages"]

    css = """
    body { font-family: 'Segoe UI', system-ui, sans-serif; max-width: 1200px; margin: 24px auto; padding: 0 16px; color: #111; }
    h1, h2, h3, h4 { color: #111; }
    code { background: #f3f4f6; padding: 2px 6px; border-radius: 3px; font-size: 90%; }
    .meta { background: #f8fafc; padding: 12px; border-radius: 6px; margin-bottom: 16px; }
    .summary { display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; margin-bottom: 16px; }
    .summary div { background: #f1f5f9; padding: 8px 10px; border-radius: 6px; font-size: 13px; }
    .summary strong { font-size: 18px; display: block; margin-top: 2px; }
    .pair { border: 1px solid #e5e7eb; border-radius: 8px; padding: 14px; margin-bottom: 16px; }
    .finding { background: #fafafa; border-left: 4px solid #94a3b8; padding: 10px 12px; margin: 8px 0; border-radius: 4px; }
    .finding.sev-high   { border-left-color: #b91c1c; }
    .finding.sev-medium { border-left-color: #b45309; }
    .finding.sev-low    { border-left-color: #1d4ed8; }
    .finding pre { background: #f8fafc; padding: 6px; border-radius: 3px; white-space: pre-wrap; font-size: 12px; }
    .finding .row { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 8px; }
    .finding .row > div { background: #fff; border: 1px solid #e5e7eb; padding: 6px; border-radius: 3px; }
    .finding img { max-width: 100%; border: 1px solid #e5e7eb; border-radius: 3px; display: block; margin-top: 4px; }
    .pill { display: inline-block; padding: 1px 6px; border-radius: 8px; font-size: 11px; font-weight: 600; }
    .pill-status { background: #e0e7ff; color: #3730a3; }
    .pill-sev { background: #fef3c7; color: #92400e; }
    .pill-type { background: #f3f4f6; color: #374151; }
    .filters { font-size: 11px; background: #f8fafc; padding: 8px; border-radius: 4px; }
    """

    def esc(x): return html.escape(str(x or ""), quote=True)

    out = [f"<!doctype html><html lang='ru'><head><meta charset='utf-8'><title>Отчёт сравнения стадий</title>"
           f"<style>{css}</style></head><body>"]
    out.append(f"<h1>Отчёт по сравнению стадий</h1>")
    out.append(f"<div class='meta'>"
               f"<div>Сессия: <code>{esc(s['id'])}</code></div>"
               f"<div>Дата формирования: {esc(data['generated_at'])}</div>"
               f"<div>Стадия A: <code>{esc(s.get('stage_a_path'))}</code></div>"
               f"<div>Стадия B: <code>{esc(s.get('stage_b_path'))}</code></div>"
               f"<div>PDF-пар: {len(data['pairs'])}</div>"
               f"<div>Листы: сопоставлено {pp['matched_pages']}, новых справа {pp['new_right_pages']}, "
               f"удалено слева {pp['removed_left_pages']}, переставлено {pp['reordered_pages']}</div>"
               f"</div>")

    warnings = data.get("warnings") or []
    if warnings:
        out.append("<h2>Предупреждения сессии</h2>")
        out.append("<ul>")
        for w in warnings:
            out.append(f"<li><strong>[{esc(w.get('severity', 'low'))}]</strong> {esc(w.get('title'))}"
                       + (f" — {esc(w.get('details'))}" if w.get('details') else "") + "</li>")
        out.append("</ul>")

    out.append("<h2>Сводка</h2>")
    out.append("<div class='summary'>")
    out.append(f"<div>Всего расхождений<strong>{su['total']}</strong></div>")
    for label, key in [("По статусам", "by_status"), ("По категориям", "by_category"),
                       ("По важности", "by_severity"), ("По типам", "by_type")]:
        items = ", ".join(f"{esc(k)}: {v}" for k, v in (su[key] or {}).items()) or "—"
        out.append(f"<div>{label}<strong style='font-size:12px;font-weight:500'>{items}</strong></div>")
    out.append("</div>")

    out.append("<h2>Расхождения по проектам</h2>")
    finding_idx = 0
    for pb in data["pairs"]:
        p = pb["pair"]
        lp = (p.get("left") or {}).get("filename", "—")
        rp = (p.get("right") or {}).get("filename", "—")
        st = pb["stats"]
        out.append(f"<div class='pair'><h3>{esc(lp)} ↔ {esc(rp)}</h3>")
        out.append(f"<p>Статус: <span class='pill pill-status'>{esc(p.get('status'))}</span> · "
                   f"листы: сопоставлено {st['matched_pages']}, новых {st['new_right_pages']}, "
                   f"удалено {st['removed_left_pages']}, переставлено {st['reordered_pages']} · "
                   f"расхождений: <strong>{len(pb['findings'])}</strong></p>")

        for it in pb["findings"]:
            finding_idx += 1
            sev_class = "sev-" + str(it.get("severity", "low"))
            out.append(f"<div class='finding {sev_class}'>")
            out.append(f"<div><strong>#{finding_idx}. {esc(it.get('title') or it.get('type'))}</strong></div>")
            out.append(f"<div style='margin-top:4px'>"
                       f"<span class='pill pill-type'>{esc(it.get('type'))}</span> "
                       f"<span class='pill pill-status'>{esc(it.get('status'))}</span> "
                       f"<span class='pill pill-sev'>{esc(it.get('severity'))}</span></div>")
            src = it.get("source") or {}
            left = it.get("left") or {}
            right = it.get("right") or {}
            if src.get("alignment_slot") is not None:
                out.append(f"<div>Slot: {src['alignment_slot']}</div>")
            if left.get("page") is not None or right.get("page") is not None:
                out.append(f"<div>L стр. {esc(left.get('page'))} / R стр. {esc(right.get('page'))}</div>")
            if it.get("summary"):
                out.append(f"<div>{esc(it['summary'])}</div>")

            row_open = False
            if left.get("text") or right.get("text") or (include_images and (left.get("crop_url") or right.get("crop_url"))):
                out.append("<div class='row'>")
                row_open = True
                # Левая колонка
                out.append("<div><strong>Стадия A</strong>")
                if left.get("text"):
                    out.append(f"<pre>{esc(str(left['text'])[:1500])}</pre>")
                if include_images and left.get("crop_url"):
                    out.append(f"<img src='{esc(left['crop_url'])}' alt='left crop' />")
                out.append("</div>")
                # Правая колонка
                out.append("<div><strong>Стадия B</strong>")
                if right.get("text"):
                    out.append(f"<pre>{esc(str(right['text'])[:1500])}</pre>")
                if include_images and right.get("crop_url"):
                    out.append(f"<img src='{esc(right['crop_url'])}' alt='right crop' />")
                out.append("</div>")
            if row_open:
                out.append("</div>")

            if include_llm and it.get("llm_summary"):
                out.append(f"<div style='margin-top:6px'><em>LLM:</em> <pre>{esc(it['llm_summary'])}</pre></div>")
            if include_user_notes and it.get("user_note"):
                out.append(f"<div style='margin-top:6px'><em>Комментарий проверяющего:</em> {esc(it['user_note'])}</div>")
            out.append("</div>")
        out.append("</div>")

    out.append("<h2>Приложение</h2>")
    out.append("<div class='filters'><strong>Параметры фильтрации:</strong><pre>")
    out.append(esc(json.dumps(data.get("filters") or {}, ensure_ascii=False, indent=2)))
    out.append(f"</pre>Findings всего в сессии (до фильтра): {data['raw_total']}</div>")
    out.append("</body></html>")
    return "".join(out)


def _render_json(session_id: str, data: dict, *, include_images: bool) -> str:
    """JSON-отчёт. При include_images=True для каждого finding с известным
    block_id добавляется local_crop_path с относительным путём внутри
    comparison/sessions/<sid>/ (системные пути не раскрываются)."""
    if include_images:
        try:
            session_root = paths_mod.session_dir(session_id).resolve()
        except Exception:
            session_root = None
        for pb in data.get("pairs") or []:
            for it in pb.get("findings") or []:
                for side in ("left", "right"):
                    p = _safe_local_path(session_id, _resolve_crop_local_path(session_id, it, side))
                    if p and session_root is not None:
                        try:
                            rel = p.relative_to(session_root)
                            (it.setdefault(side, {}))["local_crop_path"] = str(rel)
                        except ValueError:
                            pass
    return json.dumps(data, ensure_ascii=False, indent=2, default=str)


def _render_pdf(session_id: str, data: dict, out_path: Path, *, include_images: bool,
                include_llm: bool, include_user_notes: bool) -> None:
    """PDF через reportlab. Если include_images=True — вставляет реальные
    PNG-кропы из comparison/sessions/<sid>/pairs/<pid>/crops/."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle

    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]; h2 = styles["Heading2"]; h3 = styles["Heading3"]
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9, leading=12)
    small = ParagraphStyle("small", parent=body, fontSize=8, leading=10, textColor="grey")

    def esc(x): return html.escape(str(x or ""), quote=True)

    def _image_flowable(local_path: Path) -> Optional[Image]:
        try:
            from PIL import Image as PILImage
            with PILImage.open(local_path) as im:
                w, h = im.size
        except Exception:
            return None
        if w <= 0 or h <= 0:
            return None
        target_w = min(MAX_IMAGE_WIDTH_PT, w)
        target_h = h * (target_w / w)
        try:
            return Image(str(local_path), width=target_w, height=target_h)
        except Exception:
            return None

    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm)
    flow = []
    s = data["session"]; su = data["summary"]; pp = su["pages"]
    flow.append(Paragraph("Отчёт по сравнению стадий", h1))
    flow.append(Paragraph(f"Сессия: <font face='Courier'>{esc(s['id'])}</font>", body))
    flow.append(Paragraph(f"Дата: {esc(data['generated_at'])}", body))
    flow.append(Paragraph(f"Стадия A: <font face='Courier'>{esc(s.get('stage_a_path'))}</font>", body))
    flow.append(Paragraph(f"Стадия B: <font face='Courier'>{esc(s.get('stage_b_path'))}</font>", body))
    flow.append(Paragraph(
        f"PDF-пар: {len(data['pairs'])}. Листы: сопоставлено {pp['matched_pages']}, "
        f"новых {pp['new_right_pages']}, удалено {pp['removed_left_pages']}, "
        f"переставлено {pp['reordered_pages']}.", body))
    flow.append(Spacer(0, 8))

    # Предупреждения (Задача 9)
    warnings = data.get("warnings") or []
    if warnings:
        flow.append(Paragraph("Предупреждения сессии", h2))
        for w in warnings:
            flow.append(Paragraph(
                f"[{esc(w.get('severity', 'low'))}] {esc(w.get('title'))} — {esc(w.get('details', ''))}", small))
        flow.append(Spacer(0, 8))

    flow.append(Paragraph("Сводка", h2))
    flow.append(Paragraph(f"Всего расхождений: <b>{su['total']}</b>", body))
    for label, key in [("По статусам", "by_status"), ("По категориям", "by_category"),
                       ("По важности", "by_severity"), ("По типам", "by_type")]:
        items = ", ".join(f"{esc(k)}: <b>{v}</b>" for k, v in (su[key] or {}).items()) or "—"
        flow.append(Paragraph(f"{label}: {items}", body))
    flow.append(Spacer(0, 8))

    flow.append(Paragraph("Расхождения по проектам", h2))
    finding_idx = 0
    for pb in data["pairs"]:
        p = pb["pair"]
        lp = (p.get("left") or {}).get("filename", "—")
        rp = (p.get("right") or {}).get("filename", "—")
        st = pb["stats"]
        flow.append(Paragraph(f"{esc(lp)} ↔ {esc(rp)}", h3))
        flow.append(Paragraph(f"Статус: <b>{esc(p.get('status'))}</b>. "
                              f"Сопоставлено {st['matched_pages']}, новых {st['new_right_pages']}, "
                              f"удалено {st['removed_left_pages']}, переставлено {st['reordered_pages']}. "
                              f"Расхождений: <b>{len(pb['findings'])}</b>", body))
        for it in pb["findings"]:
            finding_idx += 1
            flow.append(Paragraph(f"<b>#{finding_idx}. {esc(it.get('title') or it.get('type'))}</b>", body))
            flow.append(Paragraph(
                f"Тип: <b>{esc(it.get('type'))}</b> · Статус: <b>{esc(it.get('status'))}</b> · "
                f"Важность: <b>{esc(it.get('severity'))}</b>", small))
            children = it.get("children_count") or 0
            if children:
                flow.append(Paragraph(f"На этом листе сгруппировано блоков: <b>{children}</b>", small))
            if it.get("summary"):
                flow.append(Paragraph(esc(it["summary"]), body))
            left = it.get("left") or {}; right = it.get("right") or {}
            if left.get("text"):
                flow.append(Paragraph("Стадия A:", small))
                flow.append(Paragraph(esc(str(left["text"])[:800]), body))
            if right.get("text"):
                flow.append(Paragraph("Стадия B:", small))
                flow.append(Paragraph(esc(str(right["text"])[:800]), body))
            if include_llm and it.get("llm_summary"):
                flow.append(Paragraph(f"LLM: {esc(it['llm_summary'])}", small))
            if include_user_notes and it.get("user_note"):
                flow.append(Paragraph(f"Комментарий: {esc(it['user_note'])}", small))
            if include_images and (left.get("crop_url") or right.get("crop_url")):
                left_local = _safe_local_path(session_id, _resolve_crop_local_path(session_id, it, "left"))
                right_local = _safe_local_path(session_id, _resolve_crop_local_path(session_id, it, "right"))
                left_img = _image_flowable(left_local) if left_local else None
                right_img = _image_flowable(right_local) if right_local else None
                if left_img or right_img:
                    flow.append(Spacer(0, 2))
                    cells = [
                        [
                            Paragraph("<b>Стадия A</b>", small),
                            Paragraph("<b>Стадия B</b>", small),
                        ],
                        [
                            left_img if left_img else Paragraph("Изображение недоступно", small),
                            right_img if right_img else Paragraph("Изображение недоступно", small),
                        ],
                    ]
                    t = Table(cells, colWidths=[MAX_IMAGE_WIDTH_PT + 8, MAX_IMAGE_WIDTH_PT + 8])
                    t.setStyle(TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BOX", (0, 0), (-1, -1), 0.25, "#cbd5e1"),
                        ("INNERGRID", (0, 0), (-1, -1), 0.25, "#e5e7eb"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    flow.append(t)
                else:
                    flow.append(Paragraph("Изображения недоступны для этого замечания.", small))
            flow.append(Spacer(0, 4))

    flow.append(PageBreak())
    flow.append(Paragraph("Приложение: параметры фильтрации", h2))
    flow.append(Paragraph(esc(json.dumps(data.get("filters") or {}, ensure_ascii=False)), small))
    doc.build(flow)


def _render_docx(session_id: str, data: dict, out_path: Path, *, include_images: bool,
                 include_llm: bool, include_user_notes: bool) -> None:
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    s = data["session"]; su = data["summary"]; pp = su["pages"]
    doc.add_heading("Отчёт по сравнению стадий", level=1)
    doc.add_paragraph(f"Сессия: {s['id']}")
    doc.add_paragraph(f"Дата: {data['generated_at']}")
    doc.add_paragraph(f"Стадия A: {s.get('stage_a_path')}")
    doc.add_paragraph(f"Стадия B: {s.get('stage_b_path')}")
    doc.add_paragraph(
        f"PDF-пар: {len(data['pairs'])}. Листы: сопоставлено {pp['matched_pages']}, "
        f"новых {pp['new_right_pages']}, удалено {pp['removed_left_pages']}, "
        f"переставлено {pp['reordered_pages']}."
    )

    warnings = data.get("warnings") or []
    if warnings:
        doc.add_heading("Предупреждения сессии", level=2)
        for w in warnings:
            doc.add_paragraph(f"[{w.get('severity', 'low')}] {w.get('title')} — {w.get('details', '')}")

    doc.add_heading("Сводка", level=2)
    doc.add_paragraph(f"Всего расхождений: {su['total']}")
    for label, key in [("По статусам", "by_status"), ("По категориям", "by_category"),
                       ("По важности", "by_severity"), ("По типам", "by_type")]:
        items = ", ".join(f"{k}: {v}" for k, v in (su[key] or {}).items()) or "—"
        doc.add_paragraph(f"{label}: {items}")

    doc.add_heading("Расхождения по проектам", level=2)
    finding_idx = 0
    for pb in data["pairs"]:
        p = pb["pair"]
        lp = (p.get("left") or {}).get("filename", "—")
        rp = (p.get("right") or {}).get("filename", "—")
        st = pb["stats"]
        doc.add_heading(f"{lp} ↔ {rp}", level=3)
        doc.add_paragraph(
            f"Статус: {p.get('status')}. "
            f"Сопоставлено {st['matched_pages']}, новых {st['new_right_pages']}, "
            f"удалено {st['removed_left_pages']}, переставлено {st['reordered_pages']}. "
            f"Расхождений: {len(pb['findings'])}"
        )
        for it in pb["findings"]:
            finding_idx += 1
            doc.add_paragraph(f"#{finding_idx}. {it.get('title') or it.get('type')}", style="List Number")
            meta = doc.add_paragraph()
            meta.add_run(f"Тип: {it.get('type')} · Статус: {it.get('status')} · Важность: {it.get('severity')}")
            meta.runs[0].font.size = Pt(8)
            children = it.get("children_count") or 0
            if children:
                cp = doc.add_paragraph()
                cp.add_run(f"На этом листе сгруппировано блоков: {children}").bold = True
            if it.get("summary"):
                doc.add_paragraph(str(it["summary"]))
            left = it.get("left") or {}; right = it.get("right") or {}
            if left.get("text"):
                doc.add_paragraph("Стадия A:")
                doc.add_paragraph(str(left["text"])[:800])
            if right.get("text"):
                doc.add_paragraph("Стадия B:")
                doc.add_paragraph(str(right["text"])[:800])
            if include_llm and it.get("llm_summary"):
                doc.add_paragraph(f"LLM: {it['llm_summary']}")
            if include_user_notes and it.get("user_note"):
                doc.add_paragraph(f"Комментарий проверяющего: {it['user_note']}")
            if include_images:
                left_local = _safe_local_path(session_id, _resolve_crop_local_path(session_id, it, "left"))
                right_local = _safe_local_path(session_id, _resolve_crop_local_path(session_id, it, "right"))
                if left_local:
                    doc.add_paragraph("Старая стадия:")
                    try:
                        doc.add_picture(str(left_local), width=Cm(MAX_IMAGE_WIDTH_CM))
                    except Exception:
                        doc.add_paragraph("Изображение недоступно (ошибка вставки).")
                elif left.get("crop_url"):
                    doc.add_paragraph("Изображение левой стороны недоступно.")
                if right_local:
                    doc.add_paragraph("Новая стадия:")
                    try:
                        doc.add_picture(str(right_local), width=Cm(MAX_IMAGE_WIDTH_CM))
                    except Exception:
                        doc.add_paragraph("Изображение недоступно (ошибка вставки).")
                elif right.get("crop_url"):
                    doc.add_paragraph("Изображение правой стороны недоступно.")

    doc.add_heading("Приложение: параметры фильтрации", level=2)
    doc.add_paragraph(json.dumps(data.get("filters") or {}, ensure_ascii=False, indent=2))

    doc.save(str(out_path))


# ─── Manifest (Задача 4) ─────────────────────────────────────────────────

def _manifest_path(session_id: str) -> Path:
    return paths_mod.reports_root(session_id) / "reports.json"


def _read_manifest(session_id: str) -> dict:
    p = _manifest_path(session_id)
    if not p.exists():
        return {"version": 1, "items": []}
    try:
        with open(p, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict) or not isinstance(data.get("items"), list):
            return {"version": 1, "items": []}
        return data
    except (OSError, json.JSONDecodeError):
        return {"version": 1, "items": []}


def _write_manifest(session_id: str, manifest: dict) -> None:
    p = _manifest_path(session_id)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    os.replace(tmp, p)


def _manifest_append(session_id: str, entry: dict) -> None:
    with _lock:
        m = _read_manifest(session_id)
        m["items"] = [it for it in (m.get("items") or []) if it.get("report_id") != entry.get("report_id")]
        m["items"].insert(0, entry)  # последний сверху
        _write_manifest(session_id, m)


def _manifest_lookup(session_id: str, report_id: str) -> Optional[dict]:
    m = _read_manifest(session_id)
    for it in m.get("items") or []:
        if it.get("report_id") == report_id:
            return it
    return None


def _scan_reports_dir_fallback(session_id: str) -> list[dict]:
    """Если manifest отсутствует — восстановим список из файлов
    comparison/sessions/<sid>/reports/. Поддерживает старое именование
    (comparison_report_<ts>.<fmt>) и новое (с _ms_uuid)."""
    root = paths_mod.reports_root(session_id)
    out: list[dict] = []
    for f in sorted(root.iterdir(), reverse=True):
        if not f.is_file() or f.name == "reports.json":
            continue
        rid = f.stem
        if rid.startswith("comparison_report_"):
            rid = rid[len("comparison_report_"):]
        out.append({
            "report_id": rid,
            "filename": f.name,
            "format": f.suffix.lstrip(".") or "unknown",
            "size_bytes": f.stat().st_size,
            "created_at": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "url": f"/api/stage-comparison/sessions/{session_id}/reports/{rid}/download",
        })
    return out


# ─── Public API ──────────────────────────────────────────────────────────

def create_report(
    session_id: str, fmt: str, *,
    filters: dict | None = None,
    include_rejected: bool = False,
    include_ignored: bool = False,
    include_images: bool = True,
    include_llm_summary: bool = True,
    include_user_notes: bool = True,
    include_child_findings: bool = True,
) -> dict:
    fmt = (fmt or "").lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"unsupported_format:{fmt}")

    merged_filters = dict(filters or {})
    merged_filters["include_rejected"] = include_rejected
    merged_filters["include_ignored"] = include_ignored
    merged_filters["include_child_findings"] = include_child_findings

    with _lock:
        data = _collect_report_data(session_id, merged_filters)

        report_id, fname, out_path = _allocate_filename(session_id, fmt)

        if fmt == "md":
            content = _render_markdown(data, include_images=include_images,
                                       include_llm=include_llm_summary,
                                       include_user_notes=include_user_notes)
            out_path.write_text(content, encoding="utf-8")
        elif fmt == "html":
            content = _render_html(data, include_images=include_images,
                                   include_llm=include_llm_summary,
                                   include_user_notes=include_user_notes)
            out_path.write_text(content, encoding="utf-8")
        elif fmt == "json":
            content = _render_json(session_id, data, include_images=include_images)
            out_path.write_text(content, encoding="utf-8")
        elif fmt == "pdf":
            _render_pdf(session_id, data, out_path, include_images=include_images,
                        include_llm=include_llm_summary,
                        include_user_notes=include_user_notes)
        elif fmt == "docx":
            _render_docx(session_id, data, out_path, include_images=include_images,
                         include_llm=include_llm_summary,
                         include_user_notes=include_user_notes)

        manifest_entry = {
            "report_id": report_id,
            "filename": fname,
            "format": fmt,
            "created_at": _utc_now(),
            "filters": merged_filters,
            "include_images": include_images,
            "include_llm_summary": include_llm_summary,
            "include_user_notes": include_user_notes,
            "include_child_findings": include_child_findings,
            "findings_count": data["summary"]["total"],
            "size_bytes": out_path.stat().st_size if out_path.exists() else 0,
            "url": f"/api/stage-comparison/sessions/{session_id}/reports/{report_id}/download",
        }
        _manifest_append(session_id, manifest_entry)

        return {
            "ok": True,
            "report_id": report_id,
            "filename": fname,
            "format": fmt,
            "size_bytes": manifest_entry["size_bytes"],
            "findings_included": data["summary"]["total"],
            "created_at": manifest_entry["created_at"],
            "url": manifest_entry["url"],
        }


def list_reports(session_id: str) -> list[dict]:
    """Список отчётов с метаданными. Использует manifest; при его отсутствии —
    fallback на скан reports/."""
    m = _read_manifest(session_id)
    items = list(m.get("items") or [])
    known_ids = {it.get("report_id") for it in items if it.get("report_id")}
    # Догружаем то, что есть на диске, но отсутствует в manifest (старые отчёты)
    for f in _scan_reports_dir_fallback(session_id):
        rid = f.get("report_id")
        if rid and rid not in known_ids:
            items.append(f)
            known_ids.add(rid)
    # Сортируем по created_at (по убыванию)
    items.sort(key=lambda x: x.get("created_at") or "", reverse=True)
    return items


def resolve_report_file(session_id: str, report_id: str) -> Path:
    """Безопасно вернуть путь файла отчёта (защита от path traversal).

    Сначала ищем по manifest.filename, затем — fallback: glob по
    comparison_report_<report_id>.<ext>. Старый формат с timestamp без
    suffix тоже поддерживается."""
    # 1) Manifest
    entry = _manifest_lookup(session_id, report_id)
    if entry and entry.get("filename"):
        p = paths_mod.report_path(session_id, entry["filename"])
        if p.exists() and p.is_file():
            return p

    root = paths_mod.reports_root(session_id).resolve()
    # 2) Прямое имя файла (обратная совместимость со старым list_reports)
    try:
        candidate = paths_mod.report_path(session_id, report_id)
        if candidate.exists() and candidate.is_file():
            return candidate
    except ValueError:
        pass
    # 3) Glob по report_id (новый формат: comparison_report_<rid>.<ext>)
    for ext in SUPPORTED_FORMATS:
        candidate = root / f"comparison_report_{report_id}.{ext}"
        try:
            candidate_resolved = candidate.resolve()
            candidate_resolved.relative_to(root)
        except (ValueError, OSError):
            continue
        if candidate_resolved.exists() and candidate_resolved.is_file():
            return candidate_resolved
    raise FileNotFoundError(f"report_not_found:{report_id}")


__all__ = [
    "SUPPORTED_FORMATS",
    "create_report",
    "list_reports",
    "resolve_report_file",
]
